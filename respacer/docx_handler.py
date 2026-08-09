"""
docx_handler.py
----------------
Fixes glued-together words in a .docx file WITHOUT touching formatting.

Key insight: python-docx stores text in "runs" (a contiguous span of text
sharing one style — same bold/italic/font/size/color). When a PDF->Word
converter loses spaces, it's the text INSIDE a run that's broken; the run's
formatting itself is untouched. So we only ever rewrite `run.text` in
place -- we never split runs, move runs, or touch paragraph/run styling.
This means headings stay headings, bold stays bold, tables stay tables.
"""

from docx import Document

from .segmenter import respace_text


def _fix_paragraph(paragraph, min_length, full_report):
    for run in paragraph.runs:
        if not run.text or not run.text.strip():
            continue
        fixed, changes = respace_text(run.text, min_length=min_length, return_report=True)
        if changes:
            run.text = fixed
            full_report.extend(changes)


def fix_docx(input_path: str, output_path: str, min_length: int = 12) -> list:
    """
    Read a .docx, repair glued-together words paragraph by paragraph
    (including inside tables), and save the result to `output_path`.

    Returns a report: list of {"original": ..., "fixed": ..., "confidence": ...}
    """
    doc = Document(input_path)
    report = []

    for paragraph in doc.paragraphs:
        _fix_paragraph(paragraph, min_length, report)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fix_paragraph(paragraph, min_length, report)

    # Headers / footers, if present
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                _fix_paragraph(paragraph, min_length, report)

    doc.save(output_path)
    return report
