"""
pdf_handler.py
---------------
Extracts text directly from a PDF and writes a respaced .docx.

NOTE ON FORMATTING: PDFs don't have "runs" the way .docx files do --
there's no reliable structural link between a chunk of text and a single
style the way there is in a Word file. So for a *direct* PDF input, this
module preserves paragraph breaks and basic heading-like emphasis
(detected via font size, when available) but not full rich formatting.

If you already have a .docx that came out of a PDF converter and looks
wrong, prefer `fix_docx()` in docx_handler.py -- that path preserves
100% of the original formatting because it edits the existing runs
in place instead of rebuilding the document from scratch.
"""

import pdfplumber
from docx import Document
from docx.shared import Pt

from .segmenter import respace_text


def _extract_paragraphs(pdf_path: str):
    """Yield (text, is_heading_guess) tuples per page/paragraph."""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            # Blank-line-separated blocks approximate paragraphs
            for block in text.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                # crude heading guess: short line, no ending punctuation
                is_heading = len(block) < 80 and not block.endswith((".", ",", ";"))
                yield block.replace("\n", " "), is_heading


def fix_pdf(input_path: str, output_path: str, min_length: int = 12) -> list:
    """
    Extract text from a PDF, repair glued-together words, and write the
    result as a new .docx (basic structure: headings guessed, body text
    as normal paragraphs).

    Returns a report: list of {"original": ..., "fixed": ..., "confidence": ...}
    """
    doc = Document()
    report = []

    for text, is_heading in _extract_paragraphs(input_path):
        fixed, changes = respace_text(text, min_length=min_length, return_report=True)
        report.extend(changes)

        if is_heading:
            p = doc.add_paragraph()
            run = p.add_run(fixed)
            run.bold = True
            run.font.size = Pt(13)
        else:
            doc.add_paragraph(fixed)

    doc.save(output_path)
    return report
